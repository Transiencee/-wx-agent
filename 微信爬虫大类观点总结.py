from __future__ import annotations

import datetime
import os
import random
import re
import sys
import time
from pathlib import Path

from docx import Document

# 需先在config填写api key
ROOT_PATH = Path(__file__).resolve().parent
sys.path.append(str(ROOT_PATH))

from codes.crawl_new import WXArticleDetail, WXArticleInfo
from codes.summarize import Summarizer
from codes.utils import load_config
from codes.word_proc import WordManager


# 直接把你从网页上复制的凭证粘贴到这里
WECHAT_RAW_COOKIE_STR = ""
WECHAT_TOKEN = ""
WECHAT_HEADERS = """
User-Agent: Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/143.0.0.0 Safari/537.36
"""


def gen_word_doc(word_path, paper_dict):
    word = WordManager(word_path)

    for title, content in paper_dict.items():
        word.add_paragraph(
            title,
            font_cn_name="华文中宋",
            font_size=16,
            first_line_indent=True,
            bold=True,
            line_spacing=1.5,
        )
        for para in content.split("\n"):
            word.add_paragraph(
                para,
                font_cn_name="华文中宋",
                font_size=12,
                first_line_indent=True,
                line_spacing=1.5,
            )
        word.add_paragraph(
            "\n",
            font_cn_name="华文中宋",
            font_size=12,
            first_line_indent=True,
            line_spacing=1.5,
        )

    word.close()


def batch_crawl_recent_articles(
    gzh_names_list,
    days_limit=7,
    raw_cookie_str="",
    token="",
    headers_input=None,
):
    crawler = WXArticleInfo(
        raw_cookie_str=raw_cookie_str,
        token=token,
        headers_input=headers_input,
    )
    base_storage_path = ROOT_PATH / "wechat_articles"
    base_storage_path.mkdir(parents=True, exist_ok=True)

    current_timestamp = time.time()
    cutoff_timestamp = current_timestamp - (days_limit * 24 * 60 * 60)

    print("=== 开始执行微信公众号爬取 ===")

    for gzh_name in gzh_names_list:
        print(f"\n正在搜索: {gzh_name}")

        try:
            search_res = crawler.search_gzh(gzh_name)
            if not search_res:
                print(f"未找到: {gzh_name}")
                continue

            target_gzh = search_res[0]
            fakeid = target_gzh["fakeid"]
            nickname = target_gzh["nickname"]
            print(f"锁定目标: {nickname}")

            safe_nickname = re.sub(r'[\\/*?:"<>|]', "", nickname).strip()
            current_gzh_folder = base_storage_path / safe_nickname

            time.sleep(random.randint(2, 4))

            begin = 0
            is_finished = False

            while not is_finished:
                articles = crawler.get_articles(fakeid, begin=begin, count=5)
                if not articles:
                    break

                for art in articles:
                    create_time = art["create_time"]
                    title = art["title"]
                    link = art["link"]

                    if create_time >= cutoff_timestamp:
                        date_str = datetime.datetime.fromtimestamp(create_time).strftime("%Y-%m-%d")
                        print(f"  下载: [{date_str}] {title}")

                        try:
                            detail_scraper = WXArticleDetail(
                                output_root_folder=str(current_gzh_folder),
                                article_headers=headers_input,
                            )
                            detail_scraper.scrape_wechat_article(link)
                            detail_scraper.content_to_txt()
                            detail_scraper.content_to_word()
                            print("    完成")
                            time.sleep(random.randint(3, 5))
                        except Exception as e:
                            print(f"    失败: {e}")
                    else:
                        print("  遇到旧文章，停止该公众号任务")
                        is_finished = True
                        break

                if is_finished:
                    break

                begin += 5
                time.sleep(random.randint(2, 4))

        except Exception as e:
            print(f"发生错误: {e}")


def get_local_raw_content(author_name, range_length, author_to_oa_mapping):
    folder_name = author_to_oa_mapping.get(author_name)
    if not folder_name:
        return "", 0

    oa_path = ROOT_PATH / "wechat_articles" / folder_name
    if not oa_path.exists():
        return "", 0

    cutoff_dt = datetime.datetime.now() - datetime.timedelta(days=range_length)
    valid_docs = []

    for root, _, files in os.walk(oa_path):
        for file in files:
            file_path = Path(root) / file
            if datetime.datetime.fromtimestamp(file_path.stat().st_mtime) < cutoff_dt:
                continue

            content = ""
            try:
                if file.endswith(".docx") and not file.startswith("~$"):
                    doc = Document(str(file_path))
                    content = "\n".join(
                        para.text.strip() for para in doc.paragraphs if para.text.strip()
                    )
                elif file.lower().endswith(".txt"):
                    try:
                        content = file_path.read_text(encoding="utf-8")
                    except UnicodeDecodeError:
                        content = file_path.read_text(encoding="gbk", errors="ignore")

                if content.strip():
                    valid_docs.append(f"【公众号文章：{file}】\n{content[:3000].strip()}\n")
            except Exception as e:
                print(f"  [读取失败] {file}: {e}")

    return "\n".join(valid_docs), len(valid_docs)


def build_author_prompt(asset_name, author_name, focus_point, full_context):
    return f"""
<task>
你是一位资深的{asset_name}分析师。你的任务是完全基于提供的微信公众号资料，挖掘分析师【{author_name}】关于【{asset_name}】的核心观点。
</task>

<focus_guide>
以下是需要重点关注的方向（仅供参考，不要求全中）：
{focus_point}
</focus_guide>

<requirements>
1. 宽容匹配原则：
   - 上述 focus_guide 是挖掘雷达，不是过滤网。
   - 只要资料涉及其中任何一个维度，或者讨论了与{asset_name}相关的其他重要逻辑，都必须总结出来。
   - 不要因为文章没有提到某个细分指标，就直接输出“暂无”。

2. 深度挖掘隐含观点：
   - 很多宏观文章不会直接写“{asset_name}观点”，但会给出经济、政策、流动性、风险偏好的判断。
   - 需要把这些对{asset_name}有映射关系的判断提炼出来。

3. 暂无判定：
   - 只有当资料完全与财经金融无关，或完全没有涉及{asset_name}相关逻辑时，才输出“暂无”。
   - 不要强行关联，更不要编造观点。

4. 输出风格：
   - 直接输出观点逻辑段落。
   - 语言朴素专业，不要有AI感，不要有 markdown 格式。
   - 不要写“{author_name}指出”“根据文章”等开头。
   - 严禁提及个股名称。
</requirements>

<context>
{full_context}
</context>
"""


if __name__ == "__main__":
    range_length = 7

    if not WECHAT_RAW_COOKIE_STR.strip():
        raise ValueError("请先在脚本顶部填写 WECHAT_RAW_COOKIE_STR")
    if not str(WECHAT_TOKEN).strip():
        raise ValueError("请先在脚本顶部填写 WECHAT_TOKEN")

    author_to_oa_mapping = {
        "洪灏": "洪灝的宏观策略",

        "刘刚": "Kevin策略研究",

        "张忆东": "张忆东策略世界",

    }

    target_list = list(author_to_oa_mapping.values())
    batch_crawl_recent_articles(
        target_list,
        days_limit=range_length,
        raw_cookie_str=WECHAT_RAW_COOKIE_STR,
        token=WECHAT_TOKEN,
        headers_input=WECHAT_HEADERS,
    )
    print("微信爬虫凭证已改为从当前脚本顶部读取，底层调用的是 codes/crawl_new.py。")

    asset_config = {
        "A股": {
            "authors": ["洪灏"],
            "focus": (
                "全方位研判A股市场的大势（牛熊/震荡）、风格切换（大盘vs小盘、价值vs成长、红利vs科技）及核心驱动力。"
                "1. 宏观映射: 分析GDP、通胀（CPI/PPI）、社融M1等宏观变量对企业盈利和估值的具体影响。"
                "2. 资金与情绪: 关注北向资金流向、成交量变化、两融数据及市场风险偏好。"
                "3. 配置建议: 提取分析师对具体行业板块的看法及仓位建议。"
            ),
        },
        "海外": {
            "authors": ["刘刚"],
            "focus": (
                "重点复盘与展望全球大类资产（美股、美债、大宗商品）走势。"
                "1. 美联储政策: 深度解析通胀/非农数据，预判美联储降息节奏及其对全球流动性的影响。"
                "2. 美国经济: 判断美国经济是软着陆还是衰退，以及美债利率的顶部/底部位置。"
                "3. 地缘与大选: 关注美国大选、地缘冲突对避险资产及供应链的冲击。"
            ),
        },
        "港股": {
            "authors": ["张忆东", "刘刚"],
            "focus": (
                "分析港股市场的分子端（国内基本面）与分母端（海外流动性）双重逻辑。"
                "1. 核心指标: 恒生科技/恒指的趋势判断、AH股溢价率变化、南向资金流向。"
                "2. 宏观联动: 美联储降息对港股流动性的边际改善，以及国内稳增长政策对港股盈利预期的提振。"
            ),
        },
        "黄金": {
            "authors": [],
            "focus": (
                "分析黄金作为避险资产与抗通胀资产的配置价值。"
                "1. 定价逻辑: 实际利率（美债收益率）、美元指数与金价的负相关性分析。"
                "2. 长期驱动: 全球央行购金需求、地缘政治风险带来的避险溢价，以及信用货币体系信用受损下的长期逻辑。"
            ),
        },
        "可转债": {
            "authors": [],
            "focus": (
                "分析转债市场的进攻性与防守性。"
                "1. 市场指标: 转债指数走势、平均转股溢价率、百元溢价率及纯债溢价率的变化。"
                "2. 策略判断: 正股市场对转债的驱动、债底支撑力度、下修博弈机会以及低价策略vs低溢价策略的选择。"
            ),
        },
        "国内债券": {
            "authors": [],
            "focus": (
                "研判债券牛熊方向与收益率曲线形态。"
                "1. 利率走势: 10年期/30年期国债收益率的波动区间、关键点位判断。"
                "2. 宏观驱动: 基本面预期、货币政策及债券供给对债市的冲击。"
                "3. 机构行为: 银行间流动性、理财赎回压力或配置力量分析。"
            ),
        },
        "宏观": {
            "authors": [],
            "focus": (
                "构建中国宏观经济的整体图景。"
                "1. 经济数据: 深度解读GDP、PMI、工业增加值、出口/投资/消费的边际变化。"
                "2. 通胀与金融: CPI/PPI剪刀差、社融信贷结构（M1/M2）对经济活力的指示。"
                "3. 周期位置: 判断当前处于库存周期的哪个阶段。"
            ),
        },
        "政策": {
            "authors": [],
            "focus": (
                "梳理顶层设计与政策风向标。"
                "1. 重要会议: 政治局会议、中央经济工作会议、两会/三中全会等对改革方向的定调。"
                "2. 政策组合拳: 财政、货币、产业政策的协同力度及落地效果。"
            ),
        },
        "汇率": {
            "authors": [],
            "focus": (
                "分析人民币汇率的中期趋势与短期波动因素。"
                "1. 定价因子: 美元指数强弱、中美利差、出口结汇意愿及经常账户状况。"
                "2. 央行态度: 逆周期因子、外汇存款准备金率、掉期点等汇率干预工具的使用及政策底线。"
            ),
        },
        "财政": {
            "authors": [],
            "focus": (
                "分析广义财政收支与地方债务问题。"
                "1. 财政力度: 专项债发行进度、超长期特别国债的使用、财政赤字率及广义赤字的变化。"
                "2. 隐性债务: 地方化债方案、土地出让金下滑对地方财力的掣肘以及准财政工具的应用。"
            ),
        },
        "货币": {
            "authors": [],
            "focus": (
                "判断货币政策基调（宽松/中性/收紧）。"
                "1. 量价分析: 降准降息预期、公开市场操作（OMO/MLF）的量价变化、LPR报价调整。"
                "2. 流动性: 银行间市场资金面松紧、资金利率中枢位移动，以及防空转与支持实体之间的平衡。"
            ),
        },
        "地产": {
            "authors": [],
            "focus": (
                "分析房地产市场的供需格局与政策底。"
                "1. 高频数据: 30大中城市新房/二手房成交面积、房价变动趋势、去化周期。"
                "2. 供给端: 房企拿地意愿、新开工数据、保交楼进度及房企信用风险。"
                "3. 政策端: 限购松绑、首付比例降低、房贷利率调整及收储/以旧换新政策的落地情况。"
            ),
        },
    }

    config = load_config(str(ROOT_PATH / "config.json"))
    summarizer = Summarizer(config=config, model="deepseek-v3-2-251201")

    all_asset_summaries = []

    for asset_name, config_item in asset_config.items():
        authors = config_item["authors"]
        focus_point = config_item["focus"]
        current_asset_text_list = []

        if not authors:
            all_asset_summaries.append(f"【{asset_name}】\n暂无配置分析师")
            continue

        print(f"====== 开始处理板块：{asset_name} ======")

        for author_conf in authors:
            author_name = author_conf.split("@")[0]
            print(f"  正在处理分析师：{author_name} ...")

            local_text, local_count = get_local_raw_content(
                author_name, range_length, author_to_oa_mapping
            )
            if local_text:
                print(f"    -> 公众号数据已获取 (共 {local_count} 篇)")
            else:
                print("    -> [Skip] 未获取到最近微信文章")
                current_asset_text_list.append(f"[{author_name}] 暂无")
                continue

            system_prompt = build_author_prompt(
                asset_name=asset_name,
                author_name=author_name,
                focus_point=focus_point,
                full_context=local_text,
            )

            try:
                summary = summarizer.get_model_answer(system_prompt, "请开始总结：")
                cleaned = summary.replace(asset_name, "").replace("核心观点", "").strip()
                if len(cleaned) < 5 or "暂无" in cleaned:
                    author_result = f"[{author_name}] 暂无"
                else:
                    author_result = f"[{author_name}]\n{cleaned}"
                current_asset_text_list.append(author_result)
            except Exception as e:
                print(f"    [Error] LLM生成失败: {e}")
                current_asset_text_list.append(f"[{author_name}] 生成错误")

        full_asset_block = f"【{asset_name}】\n" + "\n\n".join(current_asset_text_list)
        all_asset_summaries.append(full_asset_block)

    print("\n所有分析师处理完毕，正在生成文档...")
    final_output_text = "\n\n".join(all_asset_summaries)
    today_str = datetime.datetime.today().strftime("%Y%m%d")
    gen_word_doc(str(ROOT_PATH / f"大类观点全覆盖_仅微信_{today_str}.docx"), {"大类观点": final_output_text})
    print("完成！")
