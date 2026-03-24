from docx import Document
from docx.shared import Pt
import re
from docx.oxml.ns import  qn

class WordManager:
    
    def __init__(self, target_path):
        self.target_path = target_path
        self.document = Document()

    @staticmethod
    def split_raw_text(raw_text):
        chinese_pattern = '[\u3002\uff1b\uff0c\uff1a\u201c\u201d\uff08\uff09\u3001\uff1f\u300a\u300b\u4e00-\u9fa5]+'
        non_chinese_pattern = '[^\u3002\uff1b\uff0c\uff1a\u201c\u201d\uff08\uff09\u3001\uff1f\u300a\u300b\u4e00-\u9fa5]+'
        text_list = re.findall(non_chinese_pattern + '|' + chinese_pattern, raw_text)
        return text_list

    def add_paragraph(
        self, 
        text: str, 
        font_cn_name: str = u"楷体",
        font_en_name: str = 'Times New Roman',
        font_size: int = 11,
        first_line_indent: bool = False,
        bold: bool = False,
        space_before = Pt(0),
        space_after = Pt(0),
        line_spacing = 1# Pt(15)
    ):
        """python-docx方法格式化"""        
        p = self.document.add_paragraph()
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        p.paragraph_format.line_spacing = line_spacing
        #设置段落格式
        if first_line_indent:
            p.paragraph_format.first_line_indent = Pt(28)

        text_list = self.split_raw_text(text)
        chinese_pattern = '[\u3002\uff1b\uff0c\uff1a\u201c\u201d\uff08\uff09\u3001\uff1f\u300a\u300b\u4e00-\u9fa5]+'
        for cur_text in text_list:
            run = p.add_run(cur_text)
            if re.match(chinese_pattern, cur_text): # 当前为中文
                run.font.name = font_cn_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_cn_name)
            else:
                run.font.name = font_en_name
            run.font.size = Pt(font_size)
            run.bold = bold
        
    def close(self):
        self.document.save(self.target_path)
        
