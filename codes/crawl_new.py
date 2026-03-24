from __future__ import annotations

import os
import re
import time
import traceback
from typing import Any

import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches


def parse_cookie_str(cookie_str: str) -> dict[str, str]:
    cookie_dict: dict[str, str] = {}
    for item in cookie_str.split(";"):
        item = item.strip()
        if "=" in item:
            key, val = item.split("=", 1)
            cookie_dict[key.strip()] = val.strip()
    return cookie_dict


def parse_headers_input(headers_input: Any) -> dict[str, str]:
    if headers_input is None:
        return {}

    if isinstance(headers_input, dict):
        return {str(k).strip(): str(v).strip() for k, v in headers_input.items() if str(k).strip()}

    if isinstance(headers_input, str):
        headers: dict[str, str] = {}
        for line in headers_input.splitlines():
            line = line.strip()
            if not line or ":" not in line:
                continue
            key, value = line.split(":", 1)
            headers[key.strip()] = value.strip()
        return headers

    raise TypeError("headers_input 只支持 dict、str 或 None")


class WXArticleInfo:
    def __init__(self, raw_cookie_str: str, token: str, headers_input: Any = None):
        if not raw_cookie_str or not raw_cookie_str.strip():
            raise ValueError("raw_cookie_str 不能为空")
        if not token or not str(token).strip():
            raise ValueError("token 不能为空")

        self.cookies = parse_cookie_str(raw_cookie_str.strip())
        self.token = str(token).strip()

        default_headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/143.0.0.0 Safari/537.36"
            ),
            "Referer": (
                "https://mp.weixin.qq.com/cgi-bin/appmsg"
                f"?t=media/appmsg_edit_v2&action=edit&isNew=1&type=77&token={self.token}&lang=zh_CN"
            ),
        }
        custom_headers = parse_headers_input(headers_input)
        self.headers = {**default_headers, **custom_headers}
        self.headers["Referer"] = (
            "https://mp.weixin.qq.com/cgi-bin/appmsg"
            f"?t=media/appmsg_edit_v2&action=edit&isNew=1&type=77&token={self.token}&lang=zh_CN"
        )

    def search_gzh(self, query: str):
        params = {
            "action": "search_biz",
            "query": query,
            "begin": "0",
            "count": "3",
            "token": self.token,
            "lang": "zh_CN",
            "f": "json",
            "ajax": "1",
        }

        response = requests.get(
            "https://mp.weixin.qq.com/cgi-bin/searchbiz",
            params=params,
            cookies=self.cookies,
            headers=self.headers,
            timeout=30,
        )
        response.raise_for_status()
        data = response.json()
        if "list" not in data:
            raise ValueError(f"search_gzh 返回异常: {data}")
        return data["list"]

    def get_articles(self, fakeid: str, begin: int = 0, count: int = 4):
        params = {
            "action": "list_ex",
            "fakeid": fakeid,
            "query": "",
            "begin": str(begin),
            "count": str(count),
            "type": "9",
            "need_author_name": "1",
            "fingerprint": "cd60777421f7b8d9357a9ce1a338157b",
            "token": self.token,
            "lang": "zh_CN",
            "f": "json",
            "ajax": "1",
        }

        response = requests.get(
            "https://mp.weixin.qq.com/cgi-bin/appmsg",
            params=params,
            cookies=self.cookies,
            headers=self.headers,
            timeout=30,
        )
        response.raise_for_status()
        data = response.json()
        if "app_msg_list" not in data:
            raise ValueError(f"get_articles 返回异常: {data}")
        return data["app_msg_list"]


class WXArticleDetail:
    def __init__(self, output_root_folder="./wechat_articles/", article_headers: Any = None):
        default_headers = {
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/120.0.0.0 Safari/537.36"
            ),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
            "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
            "Accept-Encoding": "gzip, deflate",
            "Connection": "keep-alive",
        }
        custom_headers = parse_headers_input(article_headers)
        self.headers = {**default_headers, **custom_headers}
        self.output_root_folder = output_root_folder
        self.img_folder = None

        self.processed_elements = set()
        self.content_items = []
        self.img_counter = 0
        self.video_urls = []
        self.basic_info = {}

    def process_element(self, element, level=0):
        element_id = id(element)
        if element_id in self.processed_elements:
            return
        self.processed_elements.add(element_id)

        for child in element.children:
            if id(child) in self.processed_elements:
                continue

            if isinstance(child, str):
                text = child.strip()
                if not text:
                    continue

            if child.name == "img":
                img_url = child.get("data-src") or child.get("src")
                if img_url and img_url.startswith("http"):
                    try:
                        print(f"下载图片 {self.img_counter + 1}...")
                        img_response = requests.get(img_url, headers=self.headers, timeout=30)

                        if img_response.status_code == 200:
                            img_ext = ".jpg"
                            if "png" in img_url.lower():
                                img_ext = ".png"
                            elif "gif" in img_url.lower():
                                img_ext = ".gif"
                            elif "webp" in img_url.lower():
                                img_ext = ".webp"

                            img_filename = f"image_{self.img_counter + 1}{img_ext}"
                            img_path = os.path.join(self.img_folder, img_filename)

                            with open(img_path, "wb") as f:
                                f.write(img_response.content)

                            self.content_items.append(
                                {"type": "image", "path": img_path, "data": img_response.content}
                            )
                            self.img_counter += 1
                            time.sleep(0.3)
                    except Exception as e:
                        print(f"下载图片失败: {str(e)}")

                self.processed_elements.add(id(child))

            elif child.name in ["pre", "code"]:
                code_text = child.get_text().strip()
                if code_text:
                    self.content_items.append({"type": "code", "content": code_text})
                self.processed_elements.add(id(child))

            elif child.name == "blockquote":
                quote_text = child.get_text().strip()
                if quote_text:
                    self.content_items.append({"type": "quote", "content": quote_text})
                self.processed_elements.add(id(child))

            elif child.name in ["ul", "ol"]:
                list_items = []
                for li in child.find_all("li", recursive=False):
                    li_text = li.get_text().strip()
                    if li_text:
                        list_items.append(li_text)
                    self.processed_elements.add(id(li))

                if list_items:
                    self.content_items.append(
                        {"type": "list", "list_type": child.name, "items": list_items}
                    )
                self.processed_elements.add(id(child))

            elif child.name in ["p", "h1", "h2", "h3", "h4", "h5", "h6"]:
                has_img = child.find("img")

                if has_img:
                    self.process_element(child, level + 1)
                else:
                    text = child.get_text().strip()
                    if text and len(text) > 2:
                        item_type = "heading" if child.name.startswith("h") else "text"
                        self.content_items.append(
                            {
                                "type": item_type,
                                "content": text,
                                "level": child.name if child.name.startswith("h") else None,
                            }
                        )
                self.processed_elements.add(id(child))

            elif child.name in ["section", "div"]:
                self.process_element(child, level + 1)
                self.processed_elements.add(id(child))

            elif child.name == "iframe":
                video_src = child.get("data-src") or child.get("src")
                if video_src and (
                    "video" in video_src.lower()
                    or "v.qq.com" in video_src
                    or "mp.weixin.qq.com" in video_src
                ):
                    if not video_src.startswith("http"):
                        video_src = "https:" + video_src if video_src.startswith("//") else video_src

                    self.video_urls.append(video_src)
                    self.content_items.append({"type": "video", "url": video_src})
                self.processed_elements.add(id(child))

            elif child.name == "video":
                video_src = child.get("src")
                if video_src:
                    if not video_src.startswith("http"):
                        video_src = "https:" + video_src if video_src.startswith("//") else video_src

                    self.video_urls.append(video_src)
                    self.content_items.append({"type": "video", "url": video_src})
                self.processed_elements.add(id(child))

            elif isinstance(child, str):
                text = child.strip()
                if text and len(text) > 2:
                    self.content_items.append({"type": "text", "content": text})

            elif child.name:
                self.process_element(child, level + 1)

    def scrape_wechat_article(self, url):
        if not os.path.exists(self.output_root_folder):
            os.makedirs(self.output_root_folder)

        try:
            response = requests.get(url, headers=self.headers, timeout=30)
            response.encoding = "utf-8"

            if response.status_code != 200:
                return {"success": False, "error": f"HTTP请求失败，状态码: {response.status_code}"}

            soup = BeautifulSoup(response.text, "html.parser")

            title_element = soup.find("h1", {"id": "activity-name"}) or soup.find(
                "h2", {"class": "rich_media_title"}
            )
            title = title_element.get_text().strip() if title_element else "未知标题"

            date_element = soup.find("em", {"id": "publish_time"}) or soup.find(
                "span", {"class": "rich_media_meta_text"}
            )
            publish_date = date_element.get_text().strip() if date_element else time.strftime("%Y-%m-%d")

            author_element = soup.find("a", {"id": "js_name"}) or soup.find(
                "strong", {"class": "profile_nickname"}
            )
            author = author_element.get_text().strip() if author_element else "未知作者"

            safe_title = re.sub(r'[\\/*?:"<>|]', "", title)[:50]

            self.basic_info = {
                "title": title,
                "author": author,
                "publish_date": publish_date,
                "url": url,
                "safe_title": safe_title,
            }

            self.article_folder = os.path.join(self.output_root_folder, safe_title)
            if not os.path.exists(self.article_folder):
                os.makedirs(self.article_folder)

            self.img_folder = os.path.join(self.article_folder, "images")
            if not os.path.exists(self.img_folder):
                os.makedirs(self.img_folder)

            content_element = soup.find("div", {"id": "js_content"}) or soup.find(
                "div", {"class": "rich_media_content"}
            )
            if not content_element:
                return {"success": False, "error": "未找到文章正文内容"}

            self.process_element(content_element)

        except Exception as e:
            print(f"爬取失败: {str(e)}")
            traceback.print_exc()
            return {"success": False, "error": str(e)}

    def content_to_word(self):
        print("正在创建Word文档...")
        doc = Document()

        from docx.oxml.ns import qn

        style = doc.styles["Normal"]
        font = style.font
        font.name = "微软雅黑"
        font.size = None
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        heading = doc.add_heading(self.basic_info["title"], level=1)
        heading.style.font.name = "微软雅黑"
        heading.style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        para_author = doc.add_paragraph(f"作者: {self.basic_info['author']}")
        para_author.runs[0].font.name = "微软雅黑"
        para_author.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        para_date = doc.add_paragraph(f"发布日期: {self.basic_info['publish_date']}")
        para_date.runs[0].font.name = "微软雅黑"
        para_date.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        para_url = doc.add_paragraph(f"原文链接: {self.basic_info['url']}")
        para_url.runs[0].font.name = "微软雅黑"
        para_url.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

        for item in self.content_items:
            if item["type"] == "text":
                p = doc.add_paragraph(item["content"])
                for run in p.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            elif item["type"] == "heading":
                level_map = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
                level = level_map.get(item["level"], 2)
                heading_p = doc.add_heading(item["content"], level=level)
                for run in heading_p.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            elif item["type"] == "code":
                doc.add_paragraph("【代码块】")
                code_para = doc.add_paragraph(item["content"])
                code_para.style = "List Paragraph"
                for run in code_para.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            elif item["type"] == "quote":
                doc.add_paragraph("【引用】")
                quote_para = doc.add_paragraph(item["content"])
                quote_para.style = "Quote"
                for run in quote_para.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            elif item["type"] == "list":
                list_type_label = "有序列表" if item["list_type"] == "ol" else "无序列表"
                p_list_label = doc.add_paragraph(f"【{list_type_label}】")
                for run in p_list_label.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                for idx, list_item in enumerate(item["items"]):
                    if item["list_type"] == "ol":
                        p = doc.add_paragraph(f"{idx + 1}. {list_item}")
                    else:
                        p = doc.add_paragraph(f"• {list_item}")
                    for run in p.runs:
                        run.font.name = "微软雅黑"
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

            elif item["type"] == "image":
                try:
                    if not item["path"].endswith(".webp"):
                        doc.add_picture(item["path"], width=Inches(5.5))
                except Exception as e:
                    print(f"  插入图片失败: {str(e)}")

            elif item["type"] == "video":
                doc.add_paragraph("")
                video_heading = doc.add_heading("视频链接", level=2)
                for run in video_heading.runs:
                    run.font.name = "微软雅黑"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
                doc.add_paragraph(item["url"])

        doc_filename = f"{self.basic_info['safe_title']}.docx"
        doc_path = os.path.join(self.article_folder, doc_filename)
        doc.save(doc_path)
        print(f"Word文档已保存: {doc_path}")

    def content_to_txt(self):
        txt_filename = f"{self.basic_info['safe_title']}.txt"
        txt_path = os.path.join(self.article_folder, txt_filename)

        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(f"标题: {self.basic_info['title']}\n")
            f.write(f"作者: {self.basic_info['author']}\n")
            f.write(f"发布日期: {self.basic_info['publish_date']}\n")
            f.write(f"原文链接: {self.basic_info['url']}\n")
            f.write("\n\n")

            for item in self.content_items:
                if item["type"] == "text":
                    f.write(item["content"] + "\n\n")

                elif item["type"] == "heading":
                    level_map = {"h1": "#", "h2": "##", "h3": "###", "h4": "####", "h5": "#####", "h6": "######"}
                    prefix = level_map.get(item["level"], "##")
                    f.write(f"{prefix} {item['content']}\n\n")

                elif item["type"] == "code":
                    f.write("【代码块开始】\n")
                    f.write("```\n")
                    f.write(item["content"] + "\n")
                    f.write("```\n")
                    f.write("【代码块结束】\n\n")

                elif item["type"] == "quote":
                    f.write("【引用】\n")
                    for line in item["content"].split("\n"):
                        f.write(f"> {line}\n")
                    f.write("\n")

                elif item["type"] == "list":
                    list_type_label = "有序列表" if item["list_type"] == "ol" else "无序列表"
                    f.write(f"【{list_type_label}】\n")
                    for idx, list_item in enumerate(item["items"]):
                        if item["list_type"] == "ol":
                            f.write(f"  {idx + 1}. {list_item}\n")
                        else:
                            f.write(f"  • {list_item}\n")
                    f.write("\n")

                elif item["type"] == "image":
                    f.write(f"[图片: {os.path.basename(item['path'])}]\n\n")

                elif item["type"] == "video":
                    f.write("视频链接:\n")
                    f.write(f"{item['url']}\n")

        print(f"文本文件已保存: {txt_path}")
